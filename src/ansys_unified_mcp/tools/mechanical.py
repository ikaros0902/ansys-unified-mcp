#!/usr/bin/env python3
"""Mechanical MCP Server - Connects Claude to ANSYS Mechanical via gRPC."""

from __future__ import annotations
import json
import os
import sys
from datetime import datetime
from pathlib import Path
from mcp.server.fastmcp import FastMCP

_mechanical = None
_port = 0
__version__ = "1.0.0"

from ansys_unified_mcp.shared import mcp

# Use forward slashes so embedded paths in IronPython strings are safe on Windows
_OUTFILE = (Path(__file__).parent / "mech_out.txt").as_posix()
_SCRIPTFILE = (Path(__file__).parent / "mech_script.py").as_posix()


def _json(data):
    return json.dumps(data, indent=2, ensure_ascii=False)


def _esc(s: str) -> str:
    """Escape a user string for safe embedding in an IronPython double-quoted string literal."""
    return s.replace("\\", "\\\\").replace('"', '\\"')


def _check_connection():
    if _mechanical is None:
        return _json(
            {"ok": False, "error": "Not connected to Mechanical. Call connect_to_mechanical first."}
        )
    return None


def _run(script):
    """Run script in Mechanical, capturing print() output via temp files."""
    try:
        # Write the user script to a file so we avoid string-escaping issues
        with open(_SCRIPTFILE, "w", encoding="utf-8") as fh:
            fh.write(script)

        # Wrapper: redirect stdout, exec the script file, write output to disk.
        # Uses 'with' for the output file so IronPython's .NET GC closes the
        # handle before run_python_script() returns (plain open().write()
        # leaves it locked). Error text is captured so callers see it too.
        wrapper = (
            "import sys as _sys\n"
            "class _Cap:\n"
            "    def __init__(self): self.d=[]\n"
            "    def write(self,s): self.d.append(s)\n"
            "    def flush(self): pass\n"
            "_cap=_Cap()\n"
            "_orig=_sys.stdout\n"
            "_sys.stdout=_cap\n"
            "try:\n"
            "    import io as _io\n"
            '    with _io.open(r"'
            + _SCRIPTFILE.replace('"', '\\"')
            + "\", encoding='utf-8') as _sf: _code=_sf.read()\n"
            "    exec(_code, globals())\n"
            "except Exception as _e:\n"
            "    _cap.d.append('Script error: ' + str(_e) + '\\n')\n"
            "finally:\n"
            "    _sys.stdout=_orig\n"
            '    with open(r"' + _OUTFILE.replace('"', '\\"') + "\", 'w') as _f:\n"
            "        _f.write(''.join(_cap.d))\n"
        )

        _mechanical.run_python_script(wrapper)

        try:
            with open(_OUTFILE, "r", encoding="utf-8") as fh:
                result = fh.read().strip()
        except Exception:
            result = ""

        for path in (_OUTFILE, _SCRIPTFILE):
            try:
                os.remove(path)
            except Exception:
                pass

        return result if result else "(done)"
    except Exception as e:
        return "Error: " + str(e)


@mcp.tool()
def connect_to_mechanical(port: int = 10000) -> str:
    """Connect to ANSYS Mechanical via gRPC. Args: port: gRPC port"""
    global _mechanical, _port
    try:
        import ansys.mechanical.core as mech

        if _mechanical is not None:
            try:
                _mechanical.exit()
            except Exception:
                pass
            _mechanical = None
        _mechanical = mech.connect_to_mechanical(port=port)
        _port = port
        info = _run(
            "model = ExtAPI.DataModel.Project.Model\n"
            'print("Connected! Analyses: " + str(len(model.Analyses)))\n'
            "for i, a in enumerate(model.Analyses):\n"
            '    print("  [" + str(i) + "] " + str(a.Name) + " (" + str(a.AnalysisType) + ")")\n'
        )
        return _json({"ok": True, "port": port, "info": info})
    except ImportError:
        return _json({"ok": False, "error": "ansys-mechanical-core not installed."})
    except Exception as e:
        return _json({"ok": False, "error": str(e)})


@mcp.tool()
def disconnect_from_mechanical() -> str:
    """Disconnect from Mechanical."""
    global _mechanical, _port
    err = _check_connection()
    if err:
        return err
    try:
        _mechanical.exit()
    except Exception:
        pass
    _mechanical = None
    _port = 0
    return _json({"ok": True, "message": "Disconnected."})


@mcp.tool()
def check_mechanical_connection() -> str:
    """Check connection status."""
    if _mechanical is None:
        return _json({"connected": False})
    info = _run(
        "model = ExtAPI.DataModel.Project.Model\n"
        "for i, a in enumerate(model.Analyses):\n"
        '    print("[" + str(i) + "] " + str(a.Name) + " (" + str(a.AnalysisType) + ")")\n'
    )
    return _json({"connected": True, "port": _port, "info": info})


@mcp.tool()
def get_model_info() -> str:
    """Get model info: bodies, named selections, analyses."""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        "model = ExtAPI.DataModel.Project.Model\n"
        "info = {}\n"
        "bodies = []\n"
        "for body in model.Geometry.GetChildren(DataModelObjectCategory.Body, True):\n"
        '    bodies.append({"name": str(body.Name), "material": str(body.Material) if hasattr(body, "Material") else "N/A"})\n'
        'info["bodies"] = bodies\n'
        'info["named_selections"] = [str(ns.Name) for ns in model.GetChildren(DataModelObjectCategory.NamedSelection, True)]\n'
        'info["analyses"] = [{"name": str(a.Name), "type": str(a.AnalysisType)} for a in model.Analyses]\n'
        "print(json.dumps(info))\n"
    )
    try:
        return _json({"ok": True, "model_info": json.loads(result)})
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def list_materials() -> str:
    """List all materials in Engineering Data."""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        "model = ExtAPI.DataModel.Project.Model\n"
        "materials = [str(m.Name) for m in model.Materials.Children]\n"
        'print(json.dumps({"materials": materials}))\n'
    )
    try:
        return _json({"ok": True, **json.loads(result)})
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def assign_material(body_name: str, material_name: str) -> str:
    """Assign material to a body. Args: body_name, material_name"""
    err = _check_connection()
    if err:
        return err
    safe_body = _esc(body_name)
    safe_mat = _esc(material_name)
    script = (
        "model = ExtAPI.DataModel.Project.Model\n"
        "assigned = False\n"
        "for body in model.Geometry.GetChildren(DataModelObjectCategory.Body, True):\n"
        '    if str(body.Name) == "' + safe_body + '":\n'
        '        body.Material = "' + safe_mat + '"\n'
        "        assigned = True\n"
        "        break\n"
        'print("Assigned" if assigned else "Body not found: ' + safe_body + '")\n'
    )
    result = _run(script)
    return _json({"ok": "not found" not in result.lower(), "message": result})


@mcp.tool()
def set_mesh_element_size(element_size_mm: float) -> str:
    """Set global mesh element size in mm. Args: element_size_mm"""
    err = _check_connection()
    if err:
        return err
    size_m = element_size_mm / 1000.0
    script = (
        "model = ExtAPI.DataModel.Project.Model\n"
        "model.Mesh.ElementSize = Quantity(" + str(size_m) + ', "m")\n'
        'print("Element size set to ' + str(element_size_mm) + 'mm")\n'
    )
    return _json({"ok": True, "message": _run(script)})


@mcp.tool()
def generate_mesh() -> str:
    """Generate mesh. Visible live in Mechanical GUI."""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "ExtAPI.DataModel.Project.Model.Mesh.GenerateMesh()\n"
        'print("Mesh generated successfully.")\n'
    )
    return _json({"ok": True, "message": result})


@mcp.tool()
def get_mesh_statistics() -> str:
    """Get mesh node and element counts."""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        "mesh = ExtAPI.DataModel.Project.Model.Mesh\n"
        "try:\n"
        '    stats = {"nodes": int(mesh.Nodes), "elements": int(mesh.Elements)}\n'
        "except:\n"
        '    stats = {"nodes": "N/A", "elements": "N/A"}\n'
        "print(json.dumps(stats))\n"
    )
    try:
        return _json({"ok": True, "mesh_statistics": json.loads(result)})
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_fixed_support(named_selection: str, analysis_index: int = 0) -> str:
    """Add Fixed Support to a named selection. Args: named_selection, analysis_index"""
    err = _check_connection()
    if err:
        return err
    safe_ns = _esc(named_selection)
    script = (
        "import json\n"
        '_NS_NAME = "' + safe_ns + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _analysis = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "]\n"
        "    _bc = _analysis.AddFixedSupport()\n"
        "    _bc.Location = _ns\n"
        '    print(json.dumps({"ok": True, "message": "Fixed Support added"}))\n'
    )
    result = _run(script)
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_force(
    named_selection: str,
    fx_n: float = 0.0,
    fy_n: float = 0.0,
    fz_n: float = 0.0,
    analysis_index: int = 0,
) -> str:
    """Add Force load in Newtons. Args: named_selection, fx_n, fy_n, fz_n, analysis_index"""
    err = _check_connection()
    if err:
        return err
    safe_ns = _esc(named_selection)
    script = (
        "import json\n"
        '_NS_NAME = "' + safe_ns + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _analysis = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "]\n"
        "    _bc = _analysis.AddForce()\n"
        "    _bc.Location = _ns\n"
        "    _bc.DefineBy = LoadDefineBy.Components\n"
        "    _bc.XComponent.Output.SetDiscreteValue(0, Quantity(" + str(fx_n) + ', "N"))\n'
        "    _bc.YComponent.Output.SetDiscreteValue(0, Quantity(" + str(fy_n) + ', "N"))\n'
        "    _bc.ZComponent.Output.SetDiscreteValue(0, Quantity(" + str(fz_n) + ', "N"))\n'
        '    print(json.dumps({"ok": True, "message": "Force added"}))\n'
    )
    result = _run(script)
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_pressure(named_selection: str, magnitude_pa: float, analysis_index: int = 0) -> str:
    """Add Pressure load in Pascals. Args: named_selection, magnitude_pa, analysis_index"""
    err = _check_connection()
    if err:
        return err
    safe_ns = _esc(named_selection)
    script = (
        "import json\n"
        '_NS_NAME = "' + safe_ns + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _analysis = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "]\n"
        "    _bc = _analysis.AddPressure()\n"
        "    _bc.Location = _ns\n"
        "    _bc.Magnitude.Output.SetDiscreteValue(0, Quantity(" + str(magnitude_pa) + ', "Pa"))\n'
        '    print(json.dumps({"ok": True, "message": "Pressure added"}))\n'
    )
    result = _run(script)
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def list_boundary_conditions(analysis_index: int = 0) -> str:
    """List all boundary conditions. Args: analysis_index"""
    err = _check_connection()
    if err:
        return err
    script = (
        "import json\n"
        "analysis = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "]\n"
        'bcs = [{"name": str(c.Name), "type": str(c.GetType().Name)} for c in analysis.Children]\n'
        'print(json.dumps({"boundary_conditions": bcs}))\n'
    )
    result = _run(script)
    try:
        return _json({"ok": True, **json.loads(result)})
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def solve_analysis(analysis_index: int = 0) -> str:
    """Solve the analysis. Progress visible in GUI. Args: analysis_index"""
    err = _check_connection()
    if err:
        return err
    script = (
        "analysis = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "]\n"
        "analysis.Solution.Solve(True)\n"
        'print("Solve complete. Status: " + str(analysis.Solution.Status))\n'
    )
    result = _run(script)
    ok = "Script error:" not in result and "failed" not in result.lower()
    return _json({"ok": ok, "message": result})


@mcp.tool()
def get_solve_status(analysis_index: int = 0) -> str:
    """Get solve status. Args: analysis_index"""
    err = _check_connection()
    if err:
        return err
    script = (
        "import json\n"
        "analysis = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "]\n"
        'print(json.dumps({"status": str(analysis.Solution.Status), "name": str(analysis.Name), "type": str(analysis.AnalysisType)}))\n'
    )
    result = _run(script)
    try:
        return _json({"ok": True, **json.loads(result)})
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_total_deformation_all_modes(num_modes: int = 6, analysis_index: int = 0) -> str:
    """Add Total Deformation for all modal modes. Args: num_modes, analysis_index"""
    err = _check_connection()
    if err:
        return err
    script = (
        "solution = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "].Solution\n"
        "for i in range(1, " + str(num_modes + 1) + "):\n"
        "    td = solution.AddTotalDeformation()\n"
        "    td.Mode = i\n"
        "solution.EvaluateAllResults()\n"
        'print("Total Deformation added for ' + str(num_modes) + ' modes.")\n'
    )
    result = _run(script)
    return _json({"ok": True, "message": result})


@mcp.tool()
def add_total_deformation(mode: int = 0, analysis_index: int = 0) -> str:
    """Add Total Deformation result. Args: mode (0=not modal), analysis_index"""
    err = _check_connection()
    if err:
        return err
    mode_line = ("    td.Mode = " + str(mode) + "\n") if mode > 0 else ""
    script = (
        "solution = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "].Solution\n"
        "td = solution.AddTotalDeformation()\n" + mode_line + "solution.EvaluateAllResults()\n"
        'print("Total Deformation added. Max: " + str(td.Maximum))\n'
    )
    result = _run(script)
    return _json({"ok": True, "message": result})


@mcp.tool()
def get_modal_frequencies(analysis_index: int = 0) -> str:
    """Get natural frequencies from modal analysis. Args: analysis_index"""
    err = _check_connection()
    if err:
        return err
    script = (
        "import json\n"
        "solution = ExtAPI.DataModel.Project.Model.Analyses[" + str(analysis_index) + "].Solution\n"
        "frequencies = []\n"
        "for r in solution.Children:\n"
        '    if "TotalDeformation" in str(r.GetType().Name):\n'
        "        try:\n"
        '            frequencies.append({"mode": int(r.Mode), "frequency_hz": float(str(r.ReportedFrequency).split()[0])})\n'
        "        except:\n"
        "            pass\n"
        'frequencies.sort(key=lambda x: x["mode"])\n'
        'print(json.dumps({"frequencies": frequencies}))\n'
    )
    result = _run(script)
    try:
        return _json({"ok": True, **json.loads(result)})
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def generate_report(output_path: str, analysis_index: int = 0, fmt: str = "docx") -> str:
    """Generate simulation report. Args: output_path, analysis_index, fmt (docx or txt)"""
    err = _check_connection()
    if err:
        return err
    script = (
        "import json\n"
        "model = ExtAPI.DataModel.Project.Model\n"
        "analysis = model.Analyses[" + str(analysis_index) + "]\n"
        "solution = analysis.Solution\n"
        "data = {}\n"
        'data["analysis_name"] = str(analysis.Name)\n'
        'data["analysis_type"] = str(analysis.AnalysisType)\n'
        'data["solution_status"] = str(solution.Status)\n'
        'data["bodies"] = [{"name": str(b.Name), "material": str(b.Material) if hasattr(b, "Material") else "N/A"} for b in model.Geometry.GetChildren(DataModelObjectCategory.Body, True)]\n'
        "mesh = model.Mesh\n"
        "try:\n"
        '    data["nodes"] = int(mesh.Nodes)\n'
        '    data["elements"] = int(mesh.Elements)\n'
        "except:\n"
        '    data["nodes"] = "N/A"\n'
        '    data["elements"] = "N/A"\n'
        'data["boundary_conditions"] = [str(c.Name) for c in analysis.Children]\n'
        "results = []\n"
        "for r in solution.Children:\n"
        "    try:\n"
        '        rd = {"name": str(r.Name)}\n'
        '        try: rd["max"] = str(r.Maximum)\n'
        "        except: pass\n"
        '        try: rd["mode"] = str(r.Mode); rd["freq"] = str(r.ReportedFrequency)\n'
        "        except: pass\n"
        "        results.append(rd)\n"
        "    except: pass\n"
        'data["results"] = results\n'
        "print(json.dumps(data))\n"
    )
    raw = _run(script)
    try:
        data = json.loads(raw)
    except Exception:
        return _json({"ok": False, "error": "Failed to gather data", "raw": raw})
    out = Path(output_path)
    try:
        out.parent.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        return _json({"ok": False, "error": "Cannot create output directory: " + str(e)})
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    if fmt.lower() == "docx":
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            t = doc.add_heading("ANSYS Mechanical Simulation Report", 0)
            t.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("Generated: " + now)
            doc.add_paragraph(
                "Analysis: "
                + data.get("analysis_name", "N/A")
                + " ("
                + data.get("analysis_type", "N/A")
                + ")"
            )
            doc.add_paragraph("Status: " + data.get("solution_status", "N/A"))
            doc.add_heading("1. Geometry", level=1)
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            tbl.rows[0].cells[0].text = "Body"
            tbl.rows[0].cells[1].text = "Material"
            for b in data.get("bodies", []):
                row = tbl.add_row().cells
                row[0].text = b.get("name", "")
                row[1].text = b.get("material", "N/A")
            doc.add_heading("2. Mesh", level=1)
            doc.add_paragraph("Nodes: " + str(data.get("nodes", "N/A")))
            doc.add_paragraph("Elements: " + str(data.get("elements", "N/A")))
            doc.add_heading("3. Boundary Conditions", level=1)
            for bc in data.get("boundary_conditions", []):
                doc.add_paragraph(bc, style="List Bullet")
            doc.add_heading("4. Results", level=1)
            tbl2 = doc.add_table(rows=1, cols=4)
            tbl2.style = "Table Grid"
            tbl2.rows[0].cells[0].text = "Result"
            tbl2.rows[0].cells[1].text = "Mode"
            tbl2.rows[0].cells[2].text = "Frequency (Hz)"
            tbl2.rows[0].cells[3].text = "Max Value"
            for r in data.get("results", []):
                row = tbl2.add_row().cells
                row[0].text = r.get("name", "")
                row[1].text = r.get("mode", "-")
                row[2].text = r.get("freq", "-")
                row[3].text = r.get("max", "-")
            doc.save(str(out))
            return _json({"ok": True, "report_path": str(out)})
        except ImportError:
            return _json({"ok": False, "error": "python-docx not installed."})
    else:
        lines = [
            "ANSYS MECHANICAL SIMULATION REPORT",
            "Generated: " + now,
            "Analysis: "
            + data.get("analysis_name", "")
            + " ("
            + data.get("analysis_type", "")
            + ")",
            "Status: " + data.get("solution_status", ""),
            "",
            "GEOMETRY",
        ]
        for b in data.get("bodies", []):
            lines.append("  " + b.get("name", "") + " | " + b.get("material", "N/A"))
        lines += [
            "",
            "MESH",
            "  Nodes: " + str(data.get("nodes")),
            "  Elements: " + str(data.get("elements")),
            "",
            "BOUNDARY CONDITIONS",
        ]
        for bc in data.get("boundary_conditions", []):
            lines.append("  " + bc)
        lines += ["", "RESULTS"]
        for r in data.get("results", []):
            lines.append(
                "  "
                + r.get("name", "")
                + " | Mode: "
                + r.get("mode", "-")
                + " | Freq: "
                + r.get("freq", "-")
                + " | Max: "
                + r.get("max", "-")
            )
        out_txt = out.with_suffix(".txt")
        out_txt.write_text("\n".join(lines), encoding="utf-8")
        return _json({"ok": True, "report_path": str(out_txt)})


@mcp.tool()
def run_mechanical_script(script: str) -> str:
    """Run custom Python script inside Mechanical ACT API. Args: script"""
    err = _check_connection()
    if err:
        return err
    return _json({"ok": True, "output": _run(script)})


# ---------------------------------------------------------------------------
# Shared IronPython snippet: look up a named selection by name into _ns
# ---------------------------------------------------------------------------
_NS_LOOKUP = (
    "_ns = None\n"
    "for _item in ExtAPI.DataModel.Project.Model.GetChildren(DataModelObjectCategory.NamedSelection, True):\n"
    "    if str(_item.Name) == _NS_NAME:\n"
    "        _ns = _item\n"
    "        break\n"
)

# ---------------------------------------------------------------------------
# Named Selection management
# ---------------------------------------------------------------------------


@mcp.tool()
def list_named_selections() -> str:
    """List all named selections with face/body counts."""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        "out = []\n"
        "for ns in ExtAPI.DataModel.Project.Model.GetChildren(DataModelObjectCategory.NamedSelection, True):\n"
        "    try:\n"
        "        cnt = len(ns.Location.Ids) if ns.Location else 0\n"
        "    except: cnt = 0\n"
        '    out.append({"name": str(ns.Name), "entity_count": cnt})\n'
        'print(json.dumps({"named_selections": out}))\n'
    )
    try:
        return _json({"ok": True, **json.loads(result)})
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def delete_named_selection(name: str) -> str:
    """Delete a named selection by name. Args: name"""
    err = _check_connection()
    if err:
        return err
    safe_name = _esc(name)
    result = _run(
        "import json\n"
        '_NS_NAME = "' + safe_name + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Not found: ' + safe_name + '"}))\n'
        "else:\n"
        "    _ns.Delete()\n"
        '    print(json.dumps({"ok": True, "deleted": "' + safe_name + '"}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


# ---------------------------------------------------------------------------
# Geometry / body management
# ---------------------------------------------------------------------------


@mcp.tool()
def suppress_bodies(name_prefix: str = "", suppress: bool = True) -> str:
    """Suppress or unsuppress bodies by name prefix (empty = all bodies).
    Args: name_prefix, suppress (True=suppress, False=unsuppress)"""
    err = _check_connection()
    if err:
        return err
    action = "True" if suppress else "False"
    result = _run(
        "import json\n"
        '_PREFIX = "' + _esc(name_prefix) + '"\n'
        "_changed = []\n"
        "for _b in ExtAPI.DataModel.Project.Model.Geometry.GetChildren(DataModelObjectCategory.Body, True):\n"
        '    if _PREFIX == "" or str(_b.Name).startswith(_PREFIX):\n'
        "        _b.Suppressed = " + action + "\n"
        "        _changed.append(str(_b.Name))\n"
        'print(json.dumps({"ok": True, "affected": len(_changed), "bodies": _changed[:20]}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def list_point_masses() -> str:
    """List all Point Masses with CG, mass, pinball and scoped named selection."""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        "def _qm(q):\n"
        '    p=str(q).replace("[","").replace("]","").split(); v=float(p[0])\n'
        '    u=p[1] if len(p)>1 else "m"\n'
        '    if u=="mm": v/=1000.0\n'
        '    elif u=="cm": v/=100.0\n'
        "    return v\n"
        "out=[]\n"
        "for pm in ExtAPI.DataModel.Project.Model.Geometry.GetChildren(DataModelObjectCategory.PointMass, True):\n"
        "    try:\n"
        '        ns_name = str(pm.Location.Name) if pm.Location else "None"\n'
        '    except: ns_name = "None"\n'
        '    out.append({"name": str(pm.Name),\n'
        '        "mass_kg": round(_qm(pm.Mass),6),\n'
        '        "cg_mm": [round(_qm(pm.XCoordinate)*1000,1), round(_qm(pm.YCoordinate)*1000,1), round(_qm(pm.ZCoordinate)*1000,1)],\n'
        '        "pinball": str(pm.PinballRegion),\n'
        '        "named_selection": ns_name})\n'
        'print(json.dumps({"point_masses": out, "count": len(out)}))\n'
    )
    try:
        return _json({"ok": True, **json.loads(result)})
    except Exception:
        return _json({"ok": True, "raw_output": result})


# ---------------------------------------------------------------------------
# Additional boundary conditions
# ---------------------------------------------------------------------------


@mcp.tool()
def add_frictionless_support(named_selection: str, analysis_index: int = 0) -> str:
    """Add Frictionless Support. Args: named_selection, analysis_index"""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        '_NS_NAME = "' + _esc(named_selection) + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _bc = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].AddFrictionlessSupport()\n"
        "    _bc.Location = _ns\n"
        '    print(json.dumps({"ok": True, "message": "Frictionless Support added"}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_displacement(
    named_selection: str,
    x_mm: float | None = None,
    y_mm: float | None = None,
    z_mm: float | None = None,
    analysis_index: int = 0,
) -> str:
    """Add Displacement BC. Pass None for a DOF to leave it free.
    Args: named_selection, x_mm, y_mm, z_mm (None=free), analysis_index"""
    err = _check_connection()
    if err:
        return err

    def comp_line(comp, val):
        if val is None:
            return ""
        return (
            "    _bc."
            + comp
            + "Component.Output.SetDiscreteValue(0, Quantity("
            + str(val / 1000.0)
            + ', "m"))\n'
        )

    result = _run(
        "import json\n"
        '_NS_NAME = "' + _esc(named_selection) + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _bc = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].AddDisplacement()\n"
        "    _bc.Location = _ns\n"
        "    _bc.DefineBy = LoadDefineBy.Components\n"
        + comp_line("X", x_mm)
        + comp_line("Y", y_mm)
        + comp_line("Z", z_mm)
        + '    print(json.dumps({"ok": True, "message": "Displacement added"}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_remote_displacement(
    named_selection: str,
    x_mm: float | None = None,
    y_mm: float | None = None,
    z_mm: float | None = None,
    rot_x_deg: float | None = None,
    rot_y_deg: float | None = None,
    rot_z_deg: float | None = None,
    analysis_index: int = 0,
) -> str:
    """Add Remote Displacement BC. Pass None for a DOF to leave it free.
    Args: named_selection, x_mm, y_mm, z_mm, rot_x_deg, rot_y_deg, rot_z_deg, analysis_index"""
    err = _check_connection()
    if err:
        return err

    def trans_line(comp, val):
        if val is None:
            return ""
        return (
            "    _bc."
            + comp
            + "Component.Output.SetDiscreteValue(0, Quantity("
            + str(val / 1000.0)
            + ', "m"))\n'
        )

    def rot_line(comp, val):
        if val is None:
            return ""
        return (
            "    _bc."
            + comp
            + "Rotation.Output.SetDiscreteValue(0, Quantity("
            + str(val)
            + ', "deg"))\n'
        )

    result = _run(
        "import json\n"
        '_NS_NAME = "' + _esc(named_selection) + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _bc = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].AddRemoteDisplacement()\n"
        "    _bc.Location = _ns\n"
        + trans_line("X", x_mm)
        + trans_line("Y", y_mm)
        + trans_line("Z", z_mm)
        + rot_line("X", rot_x_deg)
        + rot_line("Y", rot_y_deg)
        + rot_line("Z", rot_z_deg)
        + '    print(json.dumps({"ok": True, "message": "Remote Displacement added"}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_standard_gravity(
    analysis_index: int = 0,
    x_component: float = 0.0,
    y_component: float = -1.0,
    z_component: float = 0.0,
) -> str:
    """Add Standard Earth Gravity. Direction vector defaults to -Y (down).
    Args: analysis_index, x_component, y_component, z_component (unit vector)"""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        "_grav = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].AddStandardEarthGravity()\n"
        "_grav.DefineBy = LoadDefineBy.Components\n"
        "_grav.XComponent.Output.SetDiscreteValue(0, Quantity("
        + str(x_component * 9.80665)
        + ', "m s^-2"))\n'
        "_grav.YComponent.Output.SetDiscreteValue(0, Quantity("
        + str(y_component * 9.80665)
        + ', "m s^-2"))\n'
        "_grav.ZComponent.Output.SetDiscreteValue(0, Quantity("
        + str(z_component * 9.80665)
        + ', "m s^-2"))\n'
        'print(json.dumps({"ok": True, "message": "Standard Earth Gravity added"}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_remote_force(
    named_selection: str,
    fx_n: float = 0.0,
    fy_n: float = 0.0,
    fz_n: float = 0.0,
    analysis_index: int = 0,
) -> str:
    """Add Remote Force load. Args: named_selection, fx_n, fy_n, fz_n, analysis_index"""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        '_NS_NAME = "' + _esc(named_selection) + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _bc = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].AddRemoteForce()\n"
        "    _bc.Location = _ns\n"
        "    _bc.DefineBy = LoadDefineBy.Components\n"
        "    _bc.XComponent.Output.SetDiscreteValue(0, Quantity(" + str(fx_n) + ', "N"))\n'
        "    _bc.YComponent.Output.SetDiscreteValue(0, Quantity(" + str(fy_n) + ', "N"))\n'
        "    _bc.ZComponent.Output.SetDiscreteValue(0, Quantity(" + str(fz_n) + ', "N"))\n'
        '    print(json.dumps({"ok": True, "message": "Remote Force added"}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_moment(
    named_selection: str,
    mx_nm: float = 0.0,
    my_nm: float = 0.0,
    mz_nm: float = 0.0,
    analysis_index: int = 0,
) -> str:
    """Add Moment load in N·m. Args: named_selection, mx_nm, my_nm, mz_nm, analysis_index"""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        '_NS_NAME = "' + _esc(named_selection) + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _bc = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].AddMoment()\n"
        "    _bc.Location = _ns\n"
        "    _bc.DefineBy = LoadDefineBy.Components\n"
        "    _bc.XComponent.Output.SetDiscreteValue(0, Quantity(" + str(mx_nm) + ', "N m"))\n'
        "    _bc.YComponent.Output.SetDiscreteValue(0, Quantity(" + str(my_nm) + ', "N m"))\n'
        "    _bc.ZComponent.Output.SetDiscreteValue(0, Quantity(" + str(mz_nm) + ', "N m"))\n'
        '    print(json.dumps({"ok": True, "message": "Moment added"}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


# ---------------------------------------------------------------------------
# Additional result objects
# ---------------------------------------------------------------------------


@mcp.tool()
def add_equivalent_stress(named_selection: str = "", analysis_index: int = 0) -> str:
    """Add Equivalent (von Mises) Stress result, optionally scoped to a named selection.
    Args: named_selection (empty = whole model), analysis_index"""
    err = _check_connection()
    if err:
        return err
    scope_lines = ""
    if named_selection:
        safe_ns = _esc(named_selection)
        scope_lines = '_NS_NAME = "' + safe_ns + '"\n' + _NS_LOOKUP + "if _ns: _r.Location = _ns\n"
    result = _run(
        "import json\n"
        "_solution = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].Solution\n"
        "_r = _solution.AddEquivalentStress()\n" + scope_lines + "_solution.EvaluateAllResults()\n"
        'print(json.dumps({"ok": True, "max": str(_r.Maximum), "min": str(_r.Minimum)}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_directional_deformation(
    axis: str = "Y", named_selection: str = "", analysis_index: int = 0
) -> str:
    """Add Directional Deformation result. Args: axis (X/Y/Z), named_selection, analysis_index"""
    err = _check_connection()
    if err:
        return err
    axis_map = {
        "X": "NormalOrientationType.XAxis",
        "Y": "NormalOrientationType.YAxis",
        "Z": "NormalOrientationType.ZAxis",
    }
    axis_key = axis.upper() if axis.upper() in axis_map else "Y"
    axis_enum = axis_map[axis_key]
    scope_lines = ""
    if named_selection:
        safe_ns = _esc(named_selection)
        scope_lines = '_NS_NAME = "' + safe_ns + '"\n' + _NS_LOOKUP + "if _ns: _r.Location = _ns\n"
    result = _run(
        "import json\n"
        "_solution = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].Solution\n"
        "_r = _solution.AddDirectionalDeformation()\n"
        "_r.NormalOrientation = "
        + axis_enum
        + "\n"
        + scope_lines
        + "_solution.EvaluateAllResults()\n"
        'print(json.dumps({"ok": True, "axis": "'
        + axis_key
        + '", "max": str(_r.Maximum), "min": str(_r.Minimum)}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_principal_stress(which: str = "max", analysis_index: int = 0) -> str:
    """Add Principal Stress result. Args: which (max/mid/min), analysis_index"""
    err = _check_connection()
    if err:
        return err
    method_map = {
        "max": "AddMaximumPrincipalStress",
        "mid": "AddMiddlePrincipalStress",
        "min": "AddMinimumPrincipalStress",
    }
    method = method_map.get(which.lower(), "AddMaximumPrincipalStress")
    result = _run(
        "import json\n"
        "_solution = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].Solution\n"
        "_r = _solution." + method + "()\n"
        "_solution.EvaluateAllResults()\n"
        'print(json.dumps({"ok": True, "type": "'
        + _esc(which)
        + ' principal", "max": str(_r.Maximum), "min": str(_r.Minimum)}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_stress_tool(analysis_index: int = 0) -> str:
    """Add Stress Tool (safety factor, stress ratio) using material limits. Args: analysis_index"""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        "_solution = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].Solution\n"
        "_st = _solution.AddStressTool()\n"
        "_solution.EvaluateAllResults()\n"
        'print(json.dumps({"ok": True, "message": "Stress Tool added and evaluated"}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


@mcp.tool()
def add_reaction_force(named_selection: str, analysis_index: int = 0) -> str:
    """Add Force Reaction probe at a named selection. Args: named_selection, analysis_index"""
    err = _check_connection()
    if err:
        return err
    result = _run(
        "import json\n"
        '_NS_NAME = "' + _esc(named_selection) + '"\n' + _NS_LOOKUP + "if _ns is None:\n"
        '    print(json.dumps({"ok": False, "error": "Named selection not found"}))\n'
        "else:\n"
        "    _solution = ExtAPI.DataModel.Project.Model.Analyses["
        + str(analysis_index)
        + "].Solution\n"
        "    _r = _solution.AddForceReaction()\n"
        "    _r.Location = _ns\n"
        "    _solution.EvaluateAllResults()\n"
        '    print(json.dumps({"ok": True, "x": str(_r.XAxis), "y": str(_r.YAxis), "z": str(_r.ZAxis), "total": str(_r.Total)}))\n'
    )
    try:
        return _json(json.loads(result))
    except Exception:
        return _json({"ok": True, "raw_output": result})


def _build_pm_script(match_lines: str, ns_label: str, prox_mult: float = 3.0) -> str:
    """Build the IronPython script for point-mass conversion + attachment face NS.
    match_lines: IronPython code that populates `matched` (list of body objects).
    ns_label:    suffix for the named selection name.
    No external ACT extensions required — point mass is computed from first principles.
    """
    return (
        "import math, json, binascii\n"
        "model = ExtAPI.DataModel.Project.Model\n"
        "all_bodies = model.Geometry.GetChildren(DataModelObjectCategory.Body, True)\n"
        + match_lines
        + "if not matched:\n"
        '    print(json.dumps({"ok": False, "error": "No unsuppressed bodies found"}))\n'
        "else:\n"
        "    # --- Step 1: record bounding box BEFORE suppression ---\n"
        "    eq_min = [1e18, 1e18, 1e18]\n"
        "    eq_max = [-1e18, -1e18, -1e18]\n"
        "    for _b in matched:\n"
        "        for _f in _b.GetGeoBody().Faces:\n"
        "            try:\n"
        "                _c = _f.Centroid\n"
        "                for _i in range(3):\n"
        "                    if _c[_i] < eq_min[_i]: eq_min[_i] = _c[_i]\n"
        "                    if _c[_i] > eq_max[_i]: eq_max[_i] = _c[_i]\n"
        "            except: pass\n"
        "    # --- Step 2: create Point Mass from first principles ---\n"
        "    _geometry = model.Geometry\n"
        "    _sel = ExtAPI.SelectionManager.CreateSelectionInfo(SelectionTypeEnum.GeometryEntities)\n"
        "    _sel.Ids = [_b.GetGeoBody().Id for _b in matched]\n"
        "    ExtAPI.SelectionManager.NewSelection(_sel)\n"
        "    _selection = ExtAPI.SelectionManager.CreateSelectionInfo(SelectionTypeEnum.GeometryEntities)\n"
        "    _selection.Ids = ExtAPI.SelectionManager.CurrentSelection.Ids\n"
        '    _PM_mass = Quantity("0 [kg]")\n'
        '    _PM_x    = Quantity("0 [m]")\n'
        '    _PM_y    = Quantity("0 [m]")\n'
        '    _PM_z    = Quantity("0 [m]")\n'
        '    _PM_Ixx  = Quantity("0 [kg m m]")\n'
        '    _PM_Iyy  = Quantity("0 [kg m m]")\n'
        '    _PM_Izz  = Quantity("0 [kg m m]")\n'
        "    _PM_vol  = 0.0\n"
        "    _n_bodies = 0\n"
        '    _len_unit  = "m"\n'
        '    _mass_unit = "kg"\n'
        "    _units_set = False\n"
        "    _pm = _geometry.AddPointMass()\n"
        "    ExtAPI.SelectionManager.ClearSelection()\n"
        "    for _entity in _selection.Entities:\n"
        "        _body = _geometry.GetBody(_entity)\n"
        '        if "Solid" not in str(_body.GetGeoBody().BodyType): continue\n'
        "        if not _units_set:\n"
        "            _len_unit = _body.CentroidX.Unit\n"
        "            _mass_unit = _body.Mass.Unit\n"
        "            _units_set = True\n"
        "        _bm  = _body.Mass\n"
        "        _bx  = _body.CentroidX\n"
        "        _by  = _body.CentroidY\n"
        "        _bz  = _body.CentroidZ\n"
        "        _bIx = _body.MomentOfInertiaIp1\n"
        "        _bIy = _body.MomentOfInertiaIp2\n"
        "        _bIz = _body.MomentOfInertiaIp3\n"
        '        _PM_vol += _body.Volume.ConvertUnit("m m m").Value\n'
        "        _nx = (_PM_x * _PM_mass + _bx * _bm) / (_PM_mass + _bm)\n"
        "        _ny = (_PM_y * _PM_mass + _by * _bm) / (_PM_mass + _bm)\n"
        "        _nz = (_PM_z * _PM_mass + _bz * _bm) / (_PM_mass + _bm)\n"
        "        _PM_Ixx = _PM_Ixx + _bIx + _bm*((_ny-_by)**2+(_nz-_bz)**2) + _PM_mass*((_ny-_PM_y)**2+(_nz-_PM_z)**2)\n"
        "        _PM_Iyy = _PM_Iyy + _bIy + _bm*((_nx-_bx)**2+(_nz-_bz)**2) + _PM_mass*((_nx-_PM_x)**2+(_nz-_PM_z)**2)\n"
        "        _PM_Izz = _PM_Izz + _bIz + _bm*((_nx-_bx)**2+(_ny-_by)**2) + _PM_mass*((_nx-_PM_x)**2+(_ny-_PM_y)**2)\n"
        "        _PM_mass = _PM_mass + _bm\n"
        "        _PM_x = _nx; _PM_y = _ny; _PM_z = _nz\n"
        "        _n_bodies += 1\n"
        "        _body.Suppressed = True\n"
        "    if _PM_mass.Value <= 0:\n"
        "        _pm.Delete()\n"
        '        print(json.dumps({"ok": False, "error": "No solid bodies found in selection"}))\n'
        "    else:\n"
        '        _pm_name = "%d bodies, %.5g %s" % (_n_bodies, round(_PM_mass.ConvertUnit(_mass_unit).Value, 5), _mass_unit)\n'
        "        _pinball_m = (3.0 * _PM_vol / (4.0 * math.pi)) ** (1.0 / 3.0)\n"
        "        _pm.Name = _pm_name\n"
        "        _pm.Mass = _PM_mass\n"
        "        _pm.InternalObject.LocationX = _PM_x.ConvertUnit(_len_unit).Value\n"
        "        _pm.InternalObject.LocationY = _PM_y.ConvertUnit(_len_unit).Value\n"
        "        _pm.InternalObject.LocationZ = _PM_z.ConvertUnit(_len_unit).Value\n"
        "        _pm.MassMomentOfInertiaX = _PM_Ixx\n"
        "        _pm.MassMomentOfInertiaY = _PM_Iyy\n"
        "        _pm.MassMomentOfInertiaZ = _PM_Izz\n"
        "        _pm.Activate()\n"
        "        # --- Step 3: PM CG in metres for face search ---\n"
        '        cx = _PM_x.ConvertUnit("m").Value\n'
        '        cy = _PM_y.ConvertUnit("m").Value\n'
        '        cz = _PM_z.ConvertUnit("m").Value\n'
        "        _pinball   = _pinball_m\n"
        "        _search_r  = _pinball * " + str(prox_mult) + "\n"
        "        _search_ftol = _pinball * 0.4\n"
        "        _SEED_TOL  = 0.005\n"
        "        _PLANE_TOL = 0.005\n"
        "        _FP_TOL    = 0.020\n"
        "        def _d2box_exact(c):\n"
        "            dx = max(eq_min[0]-c[0], 0.0, c[0]-eq_max[0])\n"
        "            dy = max(eq_min[1]-c[1], 0.0, c[1]-eq_max[1])\n"
        "            dz = max(eq_min[2]-c[2], 0.0, c[2]-eq_max[2])\n"
        "            return math.sqrt(dx*dx+dy*dy+dz*dz)\n"
        "        def _d2box_exp(c):\n"
        "            dx = max(eq_min[0]-_search_ftol-c[0], 0.0, c[0]-eq_max[0]-_search_ftol)\n"
        "            dy = max(eq_min[1]-_search_ftol-c[1], 0.0, c[1]-eq_max[1]-_search_ftol)\n"
        "            dz = max(eq_min[2]-_search_ftol-c[2], 0.0, c[2]-eq_max[2]-_search_ftol)\n"
        "            return math.sqrt(dx*dx+dy*dy+dz*dz)\n"
        "        # --- Step 4: seed face ON bbox boundary with best alignment to CG ---\n"
        "        _best_score = -1e18\n"
        "        _best_face  = None\n"
        "        for _body in model.Geometry.GetChildren(DataModelObjectCategory.Body, True):\n"
        "            if _body.Suppressed: continue\n"
        "            _gb = _body.GetGeoBody()\n"
        "            for _face in _gb.Faces:\n"
        "                try:\n"
        "                    c = _face.Centroid\n"
        "                    _dist_cg = math.sqrt((c[0]-cx)**2+(c[1]-cy)**2+(c[2]-cz)**2)\n"
        "                    if _dist_cg > _search_r: continue\n"
        "                    if _d2box_exp(c) > 0: continue\n"
        "                    if _d2box_exact(c) > _SEED_TOL: continue\n"
        "                    n = _face.NormalAtParam(0.5, 0.5)\n"
        "                    _dv = [cx-c[0], cy-c[1], cz-c[2]]\n"
        "                    if _dist_cg == 0: continue\n"
        "                    _align = (n[0]*_dv[0]+n[1]*_dv[1]+n[2]*_dv[2]) / _dist_cg\n"
        "                    if _align < 0.4: continue\n"
        "                    if _align > _best_score:\n"
        "                        _best_score = _align\n"
        "                        _best_face = (int(_face.Id), [c[0],c[1],c[2]], [n[0],n[1],n[2]])\n"
        "                except: pass\n"
        "        if _best_face is None:\n"
        '            print(json.dumps({"ok": False, "error": "No attachment face found on equipment boundary"}))\n'
        "        else:\n"
        "            _bf_id, _bf_c, _bf_n = _best_face\n"
        "            _bf_lvl = _bf_n[0]*_bf_c[0]+_bf_n[1]*_bf_c[1]+_bf_n[2]*_bf_c[2]\n"
        "            _abs_n  = [abs(_bf_n[0]), abs(_bf_n[1]), abs(_bf_n[2])]\n"
        "            _dom_ax = 0\n"
        "            for _i in range(1,3):\n"
        "                if _abs_n[_i] > _abs_n[_dom_ax]: _dom_ax = _i\n"
        "            _fp = [_i for _i in range(3) if _i != _dom_ax]\n"
        "            # --- Step 5: collect all coplanar faces within equipment footprint ---\n"
        "            _attach_ids = []\n"
        "            for _body in model.Geometry.GetChildren(DataModelObjectCategory.Body, True):\n"
        "                if _body.Suppressed: continue\n"
        "                _gb = _body.GetGeoBody()\n"
        "                for _face in _gb.Faces:\n"
        "                    try:\n"
        "                        c = _face.Centroid\n"
        "                        if math.sqrt((c[0]-cx)**2+(c[1]-cy)**2+(c[2]-cz)**2) > _search_r: continue\n"
        "                        if not (eq_min[_fp[0]]-_FP_TOL <= c[_fp[0]] <= eq_max[_fp[0]]+_FP_TOL): continue\n"
        "                        if not (eq_min[_fp[1]]-_FP_TOL <= c[_fp[1]] <= eq_max[_fp[1]]+_FP_TOL): continue\n"
        "                        n = _face.NormalAtParam(0.5, 0.5)\n"
        "                        if abs(n[0]*_bf_n[0]+n[1]*_bf_n[1]+n[2]*_bf_n[2]) < 0.9: continue\n"
        "                        _lvl = n[0]*c[0]+n[1]*c[1]+n[2]*c[2]\n"
        "                        if abs(_lvl - _bf_lvl) > _PLANE_TOL: continue\n"
        "                        _attach_ids.append(int(_face.Id))\n"
        "                    except: pass\n"
        "            # --- Step 6: create named selection and scope PM to it ---\n"
        "            _ns_sel = ExtAPI.SelectionManager.CreateSelectionInfo(SelectionTypeEnum.GeometryEntities)\n"
        "            _ns_sel.Ids = _attach_ids\n"
        "            _ns = model.AddNamedSelection()\n"
        '            _ns.Name = "NearestFace_' + ns_label + '"\n'
        "            _ns.Location = _ns_sel\n"
        "            _pm.Location = _ns\n"
        '            _pm.PinballRegion = Quantity("0 [mm]")\n'
        '            print(json.dumps({"ok": True,\n'
        '                "cg_mm": [round(cx*1000,1),round(cy*1000,1),round(cz*1000,1)],\n'
        '                "pinball": "0 [mm] (All)",\n'
        '                "bodies_matched": len(matched),\n'
        '                "attachment_faces": len(_attach_ids),\n'
        '                "named_selection": str(_ns.Name)}))\n'
    )


def _extract_json(raw: str):
    """Pull the first JSON object line out of raw output (past ACT log lines)."""
    for line in raw.splitlines():
        line = line.strip()
        if line.startswith("{"):
            try:
                return json.loads(line)
            except Exception:
                pass
    return {"ok": True, "raw_output": raw}


@mcp.tool()
def convert_prefix_to_point_mass(body_name_prefix: str, proximity_multiplier: float = 3.0) -> str:
    """Convert all unsuppressed bodies whose name starts with body_name_prefix to a
    combined Point Mass, then create a Named Selection of all coplanar attachment
    faces on the nearest structural surface. No ACT extensions required.
    Args: body_name_prefix, proximity_multiplier (default 3.0)"""
    err = _check_connection()
    if err:
        return err
    prefix_hex = body_name_prefix.encode("utf-8").hex()
    match = (
        "import binascii\n"
        '_PREFIX = binascii.unhexlify("' + prefix_hex + '").decode("utf-8")\n'
        "matched = [b for b in all_bodies if b.Name.startswith(_PREFIX) and not b.Suppressed]\n"
    )
    ns_label = (
        body_name_prefix[:40]
        .replace(" ", "_")
        .replace("\\", "_")
        .replace("/", "_")
        .replace('"', "_")
        .replace("\n", "")
        .replace("\r", "")
        .replace("\t", "")
    )
    raw = _run(_build_pm_script(match, ns_label, proximity_multiplier))
    return _json(_extract_json(raw))


@mcp.tool()
def convert_part_to_point_mass(part_name: str, proximity_multiplier: float = 3.0) -> str:
    """Convert all unsuppressed bodies inside a named Part to a combined Point Mass,
    then create a Named Selection of all coplanar attachment faces on the nearest
    structural surface. No ACT extensions required.
    Args: part_name, proximity_multiplier (default 3.0)"""
    err = _check_connection()
    if err:
        return err
    safe_part = _esc(part_name)
    match = (
        '_PART = "' + safe_part + '"\n'
        "matched = [b for b in all_bodies if str(b.Parent.Name) == _PART and not b.Suppressed]\n"
    )
    ns_label = (
        part_name[:40]
        .replace(" ", "_")
        .replace("\\", "_")
        .replace("/", "_")
        .replace('"', "_")
        .replace("\n", "")
        .replace("\r", "")
        .replace("\t", "")
    )
    raw = _run(_build_pm_script(match, ns_label, proximity_multiplier))
    return _json(_extract_json(raw))


def main():
    print("ANSYS Mechanical MCP Server v" + __version__, file=sys.stderr)
    print("Starting...", file=sys.stderr)
    mcp.run(transport='stdio')


if __name__ == "__main__":
    main()
